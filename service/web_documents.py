from __future__ import annotations

import io
import json
import re
import shutil
import subprocess
import tempfile
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Any, Iterable


TERM_RE = re.compile(r"^[a-z][a-z' .-]{0,119}$", re.I)

_LEGACY_DOC_NOISE = {
    "document",
    "microsoft office word",
    "microsoft word",
    "normal",
    "summaryinformation",
    "worddocument",
}


def _unique_terms(text: str) -> list[str]:
    values: list[str] = []
    seen: set[str] = set()
    for raw in re.split(r"[\r\n,，;；\t]+", text):
        term = re.sub(r"\s+", " ", raw.strip().lower())
        if term and TERM_RE.fullmatch(term) and term not in seen:
            values.append(term)
            seen.add(term)
    return values


def _legacy_doc_text(raw: bytes) -> str:
    """Recover simple word lists from legacy binary Word documents.

    Word 97-2003 files commonly keep short text runs as either single-byte
    characters or UTF-16LE.  This conservative fallback is intentionally
    limited to the English words and phrases accepted by Lexora; it avoids an
    OS package dependency while leaving formatted document conversion to
    antiword/catdoc when either is available.
    """
    candidates: list[str] = []
    # Mask UTF-16LE ranges before looking for ASCII so their alternating zero
    # bytes do not turn one word into a series of one-letter candidates.
    ascii_source = bytearray(raw)
    utf16_pattern = re.compile(
        rb"(?<![A-Za-z' .-])[A-Za-z]\x00(?:[A-Za-z' .-]\x00){0,119}"
    )
    for match in utf16_pattern.finditer(raw):
        candidates.append(match.group().decode("utf-16le", errors="ignore"))
        ascii_source[match.start():match.end()] = b"\x00" * len(match.group())
    for value in re.findall(rb"[A-Za-z][A-Za-z' .-]{0,119}", ascii_source):
        candidates.append(value.decode("ascii", errors="ignore"))

    cleaned: list[str] = []
    for candidate in candidates:
        candidate = re.sub(r"\s+", " ", candidate).strip(" .-")
        if (
            candidate
            and TERM_RE.fullmatch(candidate)
            and candidate.lower() not in _LEGACY_DOC_NOISE
        ):
            cleaned.append(candidate)
    return "\n".join(cleaned)


def extract_terms(raw: bytes, filename: str) -> list[str]:
    suffix = Path(filename).suffix.lower()
    text = ""
    if suffix in {".txt", ".text", ".md", ".csv", ".tsv"}:
        for encoding in ("utf-8-sig", "utf-16", "gb18030", "latin-1"):
            try:
                text = raw.decode(encoding)
                break
            except UnicodeDecodeError:
                continue
    elif suffix == ".pdf":
        from pypdf import PdfReader
        text = "\n".join(page.extract_text() or "" for page in PdfReader(io.BytesIO(raw)).pages)
    elif suffix == ".docx":
        from docx import Document
        document = Document(io.BytesIO(raw))
        text = "\n".join([p.text for p in document.paragraphs] + [cell.text for table in document.tables for row in table.rows for cell in row.cells])
    elif suffix == ".rtf":
        from striprtf.striprtf import rtf_to_text
        text = rtf_to_text(raw.decode("utf-8", errors="replace"))
    elif suffix == ".odt":
        from odf import teletype
        from odf.opendocument import load
        document = load(io.BytesIO(raw))
        text = teletype.extractText(document.text)
    elif suffix == ".doc":
        converter = shutil.which("antiword") or shutil.which("catdoc")
        if converter:
            source = Path(tempfile.mkstemp(prefix="lexora-import-", suffix=".doc")[1])
            try:
                source.write_bytes(raw)
                result = subprocess.run(
                    [converter, str(source)],
                    check=True,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    timeout=20,
                )
                text = result.stdout.decode("utf-8", errors="replace")
            except (subprocess.CalledProcessError, subprocess.TimeoutExpired) as error:
                raise ValueError("无法读取此 DOC 文件") from error
            finally:
                source.unlink(missing_ok=True)
        else:
            text = _legacy_doc_text(raw)
    else:
        raise ValueError("不支持此文件格式")
    terms = _unique_terms(text)
    if not terms:
        raise ValueError("文件中没有找到以换行为分隔的英语单词或短语")
    return terms


def _value(entry: dict[str, Any], key: str) -> str:
    return str(entry.get(key) or "").strip()


def _items(entry: dict[str, Any], key: str, limit: int = 8) -> list[str]:
    value = entry.get(key) or []
    if isinstance(value, str):
        try:
            value = json.loads(value)
        except json.JSONDecodeError:
            value = []
    return [str(item).strip() for item in value if str(item).strip()][:limit]


def _smart_order(entries: list[dict[str, Any]]) -> list[dict[str, Any]]:
    # Stable first-fit approximation: alternate long and short cards while
    # preserving predictable reading order inside each size group.
    scored = [(len(_value(entry, "definition")) + len(_value(entry, "definition_zh")) * 1.4 + len(_items(entry, "phrases")) * 35, index, entry) for index, entry in enumerate(entries)]
    scored.sort(key=lambda item: (-item[0], item[1]))
    result: list[dict[str, Any]] = []
    left, right = 0, len(scored) - 1
    while left <= right:
        result.append(scored[left][2])
        left += 1
        if left <= right:
            result.append(scored[right][2])
            right -= 1
    return result


def _build_docx(entries: list[dict[str, Any]], title: str, preset: str, page_size: str, typography: dict[str, float]) -> tuple[Path, str]:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_BREAK
    from docx.shared import Mm, Pt, RGBColor

    document = Document()
    section = document.sections[0]
    sizes = {"a4": (210, 297), "a5": (148, 210), "b5": (176, 250)}
    width, height = sizes.get(page_size, sizes["a4"])
    section.page_width, section.page_height = Mm(width), Mm(height)
    section.top_margin = section.bottom_margin = Mm(12)
    section.left_margin = section.right_margin = Mm(12)
    heading = document.add_heading(title or "My vocabulary book", 0)
    heading.style.font.name = "Noto Sans CJK SC"
    document.add_paragraph(f"我的双语词汇书 · {len(entries)} entries")
    columns = 3 if preset == "small" else 2 if preset == "medium" else 1
    section._sectPr.xpath("./w:cols")[0].set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num", str(columns))
    word_size = max(6, float(typography.get("word", {"small": 12, "medium": 18, "large": 24}.get(preset, 18))))
    body_size = max(6, float(typography.get("definition", {"small": 7.2, "medium": 8.7, "large": 13}.get(preset, 8.7))))
    for index, entry in enumerate(entries, 1):
        table = document.add_table(rows=1, cols=1)
        table.style = "Light Shading Accent 1"
        cell = table.cell(0, 0)
        p = cell.paragraphs[0]
        run = p.add_run(f"{index}. {_value(entry, 'word')}")
        run.bold = True
        run.font.size = Pt(word_size)
        requested = _value(entry, "requested_term")
        if requested and requested.lower() != _value(entry, "normalized_word").lower():
            match = p.add_run(f" ({requested})")
            match.font.size = Pt(max(6, word_size * .55))
            match.font.color.rgb = RGBColor(128, 134, 146)
        ipa = "  ".join(filter(None, [f"US /{_value(entry, 'us_phonetic').strip('/')}/" if _value(entry, "us_phonetic") else "", f"UK /{_value(entry, 'uk_phonetic').strip('/')}/" if _value(entry, "uk_phonetic") else ""]))
        for value, bold, blue in [
            (ipa, False, False), (_value(entry, "definition"), False, False), (_value(entry, "definition_zh"), True, True),
            ("Synonyms / 近义词  " + " · ".join(_items(entry, "synonyms")), False, False),
            ("Antonyms / 反义词  " + " · ".join(_items(entry, "antonyms")), False, False),
            ("Phrases / 常用短语  " + " · ".join(_items(entry, "phrases", 5)), False, False),
        ]:
            if not value.rstrip().endswith(("近义词", "反义词", "常用短语")) and value.strip():
                paragraph = cell.add_paragraph()
                text_run = paragraph.add_run(value)
                text_run.bold = bold
                text_run.font.size = Pt(body_size)
                if blue:
                    text_run.font.color.rgb = RGBColor(48, 76, 172)
        document.add_paragraph().paragraph_format.space_after = Pt(2)
    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    filename = f"lexora-{timestamp}.docx"
    output = Path(tempfile.mkstemp(prefix="lexora-web-", suffix=".docx")[1])
    document.save(output)
    return output, filename


def _entry_html(entry: dict[str, Any], index: int) -> str:
    import html
    esc = lambda value: html.escape(str(value or ""))
    blocks = [f"<h2><small>{index}</small> {esc(entry.get('word'))}</h2>"]
    requested = _value(entry, "requested_term")
    if requested and requested.lower() != _value(entry, "normalized_word").lower():
        blocks.append(f'<p class="requested">Original: ({esc(requested)})</p>')
    ipa = " &nbsp; ".join(filter(None, [f"US /{esc(_value(entry, 'us_phonetic').strip('/'))}/" if _value(entry, "us_phonetic") else "", f"UK /{esc(_value(entry, 'uk_phonetic').strip('/'))}/" if _value(entry, "uk_phonetic") else ""]))
    if ipa: blocks.append(f'<p class="ipa">{ipa}</p>')
    if _value(entry, "definition"): blocks.append(f"<p>{esc(entry['definition'])}</p>")
    if _value(entry, "definition_zh"): blocks.append(f'<p class="zh">{esc(entry["definition_zh"])}</p>')
    for key, label in (("synonyms", "Synonyms / 近义词"), ("antonyms", "Antonyms / 反义词"), ("phrases", "Phrases / 常用短语")):
        values = _items(entry, key)
        if values: blocks.append(f"<h3>{label}</h3><p>{esc(' · '.join(values))}</p>")
    return '<article class="entry">' + "".join(blocks) + "</article>"


def _build_epub(entries: list[dict[str, Any]], title: str, preset: str) -> tuple[Path, str]:
    import html
    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    filename = f"lexora-{timestamp}.epub"
    output = Path(tempfile.mkstemp(prefix="lexora-web-", suffix=".epub")[1])
    size = {"small": "0.82rem", "medium": "1rem", "large": "1.24rem"}.get(preset, "1rem")
    body = "".join(_entry_html(entry, index) for index, entry in enumerate(entries, 1))
    css = f"body{{font-family:system-ui,sans-serif;font-size:{size};line-height:1.55;margin:5%;color:#171923}}h1{{color:#20388f}}h2{{margin:.1em 0}}h2 small{{color:#8b92a1;font-size:.5em}}h3{{font-size:.82em;color:#304cac;margin:.7em 0 .1em}}p{{margin:.2em 0}}.entry{{break-inside:avoid;border:1px solid #dfe3ec;border-radius:12px;padding:.8em;margin:0 0 .8em;background:#f7f8fb}}.ipa,.requested{{color:#737b8c;font-size:.86em}}.zh{{color:#304cac;font-weight:600}}"
    chapter = f'<?xml version="1.0" encoding="utf-8"?><html xmlns="http://www.w3.org/1999/xhtml"><head><title>{html.escape(title)}</title><link rel="stylesheet" href="style.css" type="text/css"/></head><body><h1>{html.escape(title)}</h1><p>我的双语词汇书 · {len(entries)} entries</p>{body}</body></html>'
    identifier = f"urn:uuid:lexora-{timestamp}"
    package = f'<?xml version="1.0" encoding="utf-8"?><package xmlns="http://www.idpf.org/2007/opf" version="3.0" unique-identifier="bookid"><metadata xmlns:dc="http://purl.org/dc/elements/1.1/"><dc:identifier id="bookid">{identifier}</dc:identifier><dc:title>{html.escape(title)}</dc:title><dc:language>en</dc:language></metadata><manifest><item id="nav" href="nav.xhtml" media-type="application/xhtml+xml" properties="nav"/><item id="chapter" href="chapter.xhtml" media-type="application/xhtml+xml"/><item id="css" href="style.css" media-type="text/css"/></manifest><spine><itemref idref="chapter"/></spine></package>'
    nav = '<?xml version="1.0" encoding="utf-8"?><html xmlns="http://www.w3.org/1999/xhtml" xmlns:epub="http://www.idpf.org/2007/ops"><head><title>Contents</title></head><body><nav epub:type="toc"><ol><li><a href="chapter.xhtml">Vocabulary book</a></li></ol></nav></body></html>'
    with zipfile.ZipFile(output, "w") as archive:
        archive.writestr("mimetype", "application/epub+zip", compress_type=zipfile.ZIP_STORED)
        archive.writestr("META-INF/container.xml", '<?xml version="1.0"?><container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container"><rootfiles><rootfile full-path="OEBPS/package.opf" media-type="application/oebps-package+xml"/></rootfiles></container>')
        archive.writestr("OEBPS/package.opf", package)
        archive.writestr("OEBPS/nav.xhtml", nav)
        archive.writestr("OEBPS/chapter.xhtml", chapter)
        archive.writestr("OEBPS/style.css", css)
    return output, filename


def _wrapped(draw: Any, text: str, font: Any, width: int) -> list[str]:
    # Pillow cannot measure text containing line breaks. Dictionary entries can
    # contain both real newlines and escaped ``\\n`` separators, so collapse
    # both forms before applying the width-aware wrapping algorithm.
    text = re.sub(r"\s+", " ", str(text).replace("\\n", " ")).strip()
    words = re.findall(r"\S+\s*", text) if re.search(r"[A-Za-z]", text) else list(text)
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = current + word
        if current and draw.textlength(candidate, font=font) > width:
            lines.append(current.rstrip())
            current = word.lstrip()
        else:
            current = candidate
    if current.strip():
        lines.append(current.rstrip())
    return lines


def _build_images(entries: list[dict[str, Any]], title: str, preset: str, page_size: str, long_image: bool, typography: dict[str, float]) -> tuple[Path, str, str]:
    from PIL import Image, ImageDraw, ImageFont
    from service.web_pdf import FONT_BOLD, FONT_IPA, FONT_REGULAR
    sizes = {"a4": (1240, 1754), "a5": (874, 1240), "b5": (1039, 1476)}
    page_width, page_height = sizes.get(page_size, sizes["a4"])
    margin, gap, header = 48, 18, 92
    columns = 3 if preset == "small" else 2 if preset == "medium" else 1
    column_width = (page_width - margin * 2 - gap * (columns - 1)) // columns
    scale = 2.1
    word_font = ImageFont.truetype(str(FONT_BOLD), int(max(6, typography.get("word", 18)) * scale))
    body_font = ImageFont.truetype(str(FONT_REGULAR), int(max(6, typography.get("definition", 8.7)) * scale))
    zh_font = ImageFont.truetype(str(FONT_BOLD), int(max(6, typography.get("definition", 8.7)) * scale))
    ipa_font = ImageFont.truetype(str(FONT_IPA), int(max(6, typography.get("phonetic", 9)) * scale))
    label_font = ImageFont.truetype(str(FONT_BOLD), int(max(6, typography.get("related", 7.2)) * scale))

    probe = Image.new("RGB", (page_width, page_height), "white")
    draw = ImageDraw.Draw(probe)
    cards: list[tuple[int, list[tuple[str, Any, str]]]] = []
    for index, entry in enumerate(entries, 1):
        lines: list[tuple[str, Any, str]] = [(f"{index}. {_value(entry, 'word')}", word_font, "#11131b")]
        ipa = "  ".join(filter(None, [f"US /{_value(entry, 'us_phonetic').strip('/')}/" if _value(entry, "us_phonetic") else "", f"UK /{_value(entry, 'uk_phonetic').strip('/')}/" if _value(entry, "uk_phonetic") else ""]))
        if ipa: lines.append((ipa, ipa_font, "#707787"))
        for text, font, color in ((_value(entry, "definition"), body_font, "#242936"), (_value(entry, "definition_zh"), zh_font, "#304cac")):
            lines.extend((line, font, color) for line in _wrapped(draw, text, font, column_width - 34)[:5])
        for key, label in (("synonyms", "Synonyms / 近义词"), ("antonyms", "Antonyms / 反义词"), ("phrases", "Phrases / 常用短语")):
            values = _items(entry, key, 8)
            if values:
                lines.append((label, label_font, "#304cac"))
                lines.extend((line, body_font, "#242936") for line in _wrapped(draw, " · ".join(values), body_font, column_width - 34)[:4])
        line_height = sum(font.getbbox(text or "Ag")[3] - font.getbbox(text or "Ag")[1] + 7 for text, font, _ in lines)
        cards.append((line_height + 28, lines))

    page_cards: list[list[tuple[int, int, int, list[tuple[str, Any, str]]]]] = [[]]
    heights = [header] * columns
    for height, lines in cards:
        column = min(range(columns), key=lambda value: heights[value])
        if heights[column] + height > page_height - margin:
            page_cards.append([]); heights = [header] * columns; column = 0
        x = margin + column * (column_width + gap)
        y = heights[column]
        page_cards[-1].append((x, y, height, lines)); heights[column] += height + gap

    images: list[Image.Image] = []
    for page_index, placements in enumerate(page_cards, 1):
        image = Image.new("RGB", (page_width, page_height), "white")
        canvas = ImageDraw.Draw(image)
        canvas.text((margin, 31), "LEXORA", font=word_font, fill="#20388f")
        canvas.text((page_width - margin, 40), f"{page_index} / {len(page_cards)}", font=body_font, fill="#8b92a1", anchor="ra")
        for x, y, height, lines in placements:
            canvas.rounded_rectangle((x, y, x + column_width, y + height), radius=17, fill="#f7f8fb", outline="#dfe3ec", width=1)
            cursor = y + 14
            for text, font, color in lines:
                canvas.text((x + 16, cursor), text, font=font, fill=color)
                box = font.getbbox(text or "Ag"); cursor += box[3] - box[1] + 7
        images.append(image)
    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    if long_image:
        width, height = page_width, sum(image.height for image in images)
        canvas = Image.new("RGB", (width, height), "white")
        y = 0
        for image in images:
            canvas.paste(image, (0, y)); y += image.height
        output = Path(tempfile.mkstemp(prefix="lexora-web-", suffix=".jpg")[1])
        canvas.save(output, quality=92, optimize=True)
        return output, f"lexora-{timestamp}-long.jpg", "image/jpeg"
    output = Path(tempfile.mkstemp(prefix="lexora-web-", suffix=".zip")[1])
    with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as archive:
        for index, image in enumerate(images, 1):
            buffer = io.BytesIO(); image.save(buffer, "PNG", optimize=True)
            archive.writestr(f"lexora-page-{index:03}.png", buffer.getvalue())
    return output, f"lexora-{timestamp}-images.zip", "application/zip"


def build_document(entries: Iterable[dict[str, Any]], title: str, preset: str, example_count: int, output_format: str, page_size: str, smart_reorder: bool, typography: dict[str, float]) -> tuple[Path, str, str]:
    data = list(entries)
    if smart_reorder:
        data = _smart_order(data)
    if output_format == "docx":
        path, filename = _build_docx(data, title, preset, page_size, typography)
        return path, filename, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if output_format == "epub":
        path, filename = _build_epub(data, title, preset)
        return path, filename, "application/epub+zip"
    if output_format in {"images", "longImage"}:
        return _build_images(data, title, preset, page_size, output_format == "longImage", typography)
    from service.web_pdf import build_pdf
    pdf_path, pdf_filename = build_pdf(data, title, preset, example_count, page_size, typography)
    return pdf_path, pdf_filename, "application/pdf"
